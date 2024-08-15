from docx import Document

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('Document Title', 0)

# Add a paragraph
doc.add_paragraph('This is a paragraph in the document.')

# Add another paragraph with bold text
bold_paragraph = doc.add_paragraph('This is another paragraph with some bold text.')
bold_paragraph.runs[0].bold = True

# Add a table
table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = 'Cell 1,1'
table.cell(0, 1).text = 'Cell 1,2'
table.cell(1, 0).text = 'Cell 2,1'
table.cell(1, 1).text = 'Cell 2,2'

# Save the document
file_path = 'example1.docx'
doc.save(file_path)

print(file_path)
