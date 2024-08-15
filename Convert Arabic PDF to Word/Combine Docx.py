import os

import docx
from docx import Document
from docxcompose.composer import Composer

def add_page_break(doc):
    """
    Adds a page break to the end of the document.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

def combine_docx_files(folder_path, output_file):
    # Get a list of all DOCX files in the folder
    docx_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    docx_files.sort()  # Optional: sort files if needed

    # Create a new Document for the output
    combined_document = Document()

    # Use Composer to merge documents
    composer = Composer(combined_document)

    for index, docx_file in enumerate(docx_files):
        doc_path = os.path.join(folder_path, docx_file)
        doc = Document(doc_path)

        # Append the document
        composer.append(doc)

        # Add a page break if it's not the last document
        if index < len(docx_files) - 1:
            add_page_break(combined_document)

    # Save the combined document
    composer.save(output_file)

# Specify the folder path containing the DOCX files and the output file name
folder_path = 'C:\\Users\\Samir\\PycharmProjects\\email\\OutPutDocx'
output_file = 'C:\\Users\\Samir\\PycharmProjects\\\email\\FinalDocx\\Final.docx'


combine_docx_files(folder_path, output_file)
